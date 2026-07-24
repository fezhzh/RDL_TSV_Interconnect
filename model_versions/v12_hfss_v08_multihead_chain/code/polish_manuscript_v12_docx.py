from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


VERSION_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = VERSION_DIR / "Manuscript_v12.docx"


def set_run_font(run, source_run=None) -> None:
    if source_run is not None:
        run.bold = source_run.bold
        run.italic = source_run.italic
        run.underline = source_run.underline
        run.font.size = source_run.font.size
        run.font.name = source_run.font.name
    if run.font.name is None:
        run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), run.font.name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), run.font.name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "SimSun")


def replace_text(paragraph, text: str) -> None:
    source = paragraph.runs[0] if paragraph.runs else None
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    set_run_font(run, source)


def replace_by_index(doc: Document, index: int, text: str) -> None:
    replace_text(doc.paragraphs[index], text)


def main() -> None:
    doc = Document(DOCX_PATH)

    replacements = {
        0: "HFSS-Based Compact Modeling of Cascaded RDL-TSV Interconnects Using Equivalent Circuits and Transition Correction",
        9: (
            "Abstract-This paper presents an HFSS-based compact modeling method for cascaded "
            "redistribution-layer (RDL) and through-silicon-via (TSV) interconnect structures. "
            "Scalable equivalent-circuit models are first built for individual RDL and TSV devices, "
            "and an RDL-TSV Transition Model is inserted at each junction to capture local parasitic "
            "effects. To reduce parameter ambiguity in long cascades, shared seven-parameter "
            "optimization is followed by a twelve-head neural network for position-dependent "
            "correction. For the TSV-RDL3 validation case, the model achieves 3.1% average NMSE "
            "using 120 training samples and 58 test samples. The resulting model is compact, "
            "interpretable, SPICE-compatible, and suitable for PDK-oriented applications."
        ),
        10: "Keywords-RDL, TSV, equivalent circuit, compact model, transition model",
        12: (
            "Three-dimensional (3-D) integration enables high-density heterogeneous systems by "
            "vertically stacking dies and connecting them through redistribution layers (RDLs) "
            "and through-silicon vias (TSVs). As RDL-TSV interconnect chains become longer, their "
            "frequency response is no longer determined by isolated RDL or TSV sections alone. "
            "Transition discontinuities, repeated cascade errors, and phase accumulation jointly "
            "affect insertion loss, return loss, and delay."
        ),
        13: (
            "Full-wave electromagnetic simulation can accurately characterize these effects, but "
            "directly simulating every full-chain geometry in HFSS is too expensive for design-space "
            "exploration, compact-model library construction, and PDK-oriented development. An "
            "accurate compact model with clear circuit meaning and SPICE compatibility is therefore "
            "needed."
        ),
        14: (
            "Existing frequency-domain models usually follow two routes. Direct neural-network "
            "S-parameter surrogates can fit sampled responses, but they have limited physical "
            "interpretability and are difficult to reuse in circuit simulation. Cascading compact "
            "models of individual devices is more interpretable, but it ignores parasitic effects "
            "introduced at RDL/TSV transitions. For long chains, these small interface errors can "
            "accumulate into significant full-structure deviations."
        ),
        15: (
            "This work proposes a compact modeling flow that combines HFSS-derived equivalent-circuit "
            "models for individual RDL and TSV sections with a seven-parameter pi-type RDL-TSV "
            "Transition Model. A shared-to-multi-head neural network maps device geometry to "
            "transition-model parameters for the twelve junctions in the complete chain. The resulting "
            "model keeps the physical interpretability of equivalent circuits while improving "
            "full-chain accuracy over a direct cascade of isolated single-device models."
        ),
        18: (
            "Fig. 1 summarizes the overall modeling flow for the cascaded RDL-TSV structure. First, "
            "equivalent-circuit models are established for individual RDL and TSV devices. Based on "
            "the selected circuit topologies, the element values are extracted from HFSS-simulated "
            "S-parameters. For each circuit element, an MLP maps the physical device parameters to "
            "the corresponding element value, producing scalable compact models for individual RDL "
            "and TSV sections."
        ),
        19: (
            "After the scalable RDL and TSV models are constructed, they are used to replace the "
            "corresponding sections in the complete interconnect structure. At each RDL/TSV junction, "
            "a pi-type transition model is inserted to characterize coupling, discontinuity, and loss "
            "caused by the physical transition between the two device types. The complete compact "
            "model is then formed by cascading multiple RDL models, TSV models, and RDL-TSV "
            "Transition Models according to the physical order of the interconnect chain."
        ),
        23: (
            "For a chain with many RDL/TSV junctions, independently optimizing all transition models "
            "introduces many circuit parameters and can lead to non-unique solutions. Therefore, a "
            "two-step training strategy is used. The transition model is first optimized in a shared "
            "form, where the same circuit parameters are repeated at all connection positions. The "
            "shared optimized targets are used to train an initial parameter network. The network is "
            "then expanded into a multi-head model with twelve output heads and trained against the "
            "full-chain S-parameter response. This strategy captures position-dependent interface "
            "behavior while reducing parameter ambiguity."
        ),
        24: "RDL Equivalent-Circuit Model and Parameter Extraction Method",
        25: (
            "The RDL model is built from the equivalent-circuit topology reported in [2], as shown in "
            "Fig. 2. This circuit accurately represents the frequency response of an RDL section, and "
            "its element values can be extracted from S-parameters. For each HFSS sample, the RDL "
            "S-parameters are converted into an ABCD matrix. The characteristic impedance Zc and "
            "propagation constant gamma are then calculated as"
        ),
        26: "Using Zc and gamma, the per-unit-length RLGC curves are obtained as",
        29: (
            "Here, omega1 is selected at the lowest frequency, omega3 is selected at the highest "
            "frequency, and omega2 is selected at an intermediate frequency. In this work, these "
            "angular frequencies are set to 0.1 x 10^9 x 2pi rad/s, 20 x 10^9 x 2pi rad/s, and "
            "100 x 10^9 x 2pi rad/s, respectively."
        ),
        30: "TSV Equivalent-Circuit Model and Parameter Extraction Method",
        31: (
            "The TSV model is based on the equivalent-circuit model proposed in [3], as shown in "
            "Fig. 3. This circuit includes the TSV conductor resistance R_TSV and inductance L_TSV, "
            "oxide capacitance C_ox, silicon-substrate capacitance C_si, and substrate loss resistance "
            "R_si. To simplify parameter extraction, the original equivalent circuit is transformed "
            "into the alternative circuit shown in Fig. 4. Although the two circuits have different "
            "topologies, they provide the same frequency response and are therefore equivalent for "
            "S-parameter modeling."
        ),
        32: (
            "Similar to the RDL extraction, the TSV S-parameters are converted into ABCD matrices and "
            "used to derive RLGC curves. The parameters C_ox, C_si, and R_si are extracted from the "
            "G(omega) and C(omega) curves using the same procedure as the RDL shunt-branch extraction "
            "in (7)-(9). R_TSV is taken from the low-frequency value of R(omega), while L_TSV is taken "
            "from the high-frequency value of L(omega)."
        ),
        35: (
            "The proposed RDL-TSV Transition Model is shown in Fig. 5. This model is inserted at each "
            "RDL/TSV junction to compensate for parasitic effects that are not captured by directly "
            "cascading isolated RDL and TSV compact models. It is implemented as a seven-parameter "
            "pi-type equivalent circuit. The shunt capacitance-resistance branches model local "
            "capacitive coupling and dielectric or substrate loss near both sides of the transition, "
            "while the series resistance and inductance describe additional transition loss, current "
            "crowding, and local current-path inductance introduced by the cascade discontinuity."
        ),
        36: "Validation and Discussion",
        38: (
            "For the RDL and TSV single-device models, 400 geometry samples are generated in their "
            "parameter spaces using Latin hypercube sampling, and their S-parameters are obtained by "
            "HFSS simulation. The equivalent-circuit parameters are extracted using the models and "
            "formulas described in Sections II-B and II-C. For each circuit element parameter, an MLP "
            "with two 20-neuron hidden layers is trained to map the physical device parameters to the "
            "corresponding element value. The trained RDL and TSV models achieve average NMSE values "
            "of 0.04% and 0.02%, respectively, indicating that the single-device compact models can "
            "accurately reproduce the HFSS responses."
        ),
        40: (
            "After the individual RDL and TSV device models are established, these compact models are "
            "used to replace the corresponding RDL and TSV sections in the complete structure. The "
            "replaced device sections are then connected by the RDL-TSV Transition Model, forming a "
            "full-chain circuit composed of RDL models, transition models, and TSV models."
        ),
        41: (
            "For the RDL-TSV Transition Model, 120 samples are used for training and 58 samples are "
            "used for testing. The initial parameter network uses a 30 x 30 x 20 MLP structure. The "
            "final model is expanded into a multi-head network with twelve output heads, where each "
            "head predicts the transition-model parameters for one RDL/TSV junction. Using the "
            "S-parameters of the complete cascaded structure as the target, the final full-chain model "
            "achieves an average NMSE of 3.1% on the test set."
        ),
        43: (
            "The final average NMSE of 3.1% is slightly better than the accuracy reported in [2] "
            "using the same number of training samples."
        ),
        44: (
            "The proposed model has two practical advantages. First, it is more compact than a direct "
            "full-chain S-parameter neural surrogate because it learns only the parameters of "
            "physically meaningful equivalent circuits. Second, the resulting model can be implemented "
            "in a SPICE-compatible form, making it suitable for compact-model library construction, "
            "circuit-level verification, and PDK-oriented development."
        ),
        46: (
            "This work proposes a compact modeling method for cascaded RDL-TSV structures. By combining "
            "scalable RDL and TSV single-device models with an RDL-TSV Transition Model, the method "
            "builds an accurate high-dimensional compact model using limited full-wave data. A "
            "shared-to-multi-head training strategy reduces parasitic-parameter ambiguity while "
            "capturing position-dependent transition effects. For the TSV-RDL3 case, the model "
            "achieves 3.1% average NMSE with 120 training and 58 test samples. Since the final model "
            "is based on equivalent circuits, it more naturally preserves passivity and causality, "
            "supports SPICE simulation, and is suitable for PDK-oriented applications."
        ),
        48: (
            "[1] Z.-X. Ye et al., \"Frequency-Domain Modeling of Interconnects Based on Assemble "
            "Neural Network for 3-D Integration,\" IEEE Transactions on Computer-Aided Design of "
            "Integrated Circuits and Systems, vol. 45, no. 5, pp. 2208-2221, 2026."
        ),
        49: (
            "[2] H.-S. Yin et al., \"A SPICE-Compatible Model for On-Chip Coplanar Coupled Waveguides "
            "Up to 110 GHz,\" in 2024 IEEE 7th International Conference on Electronic Information "
            "and Communication Technology (ICEICT), Xi'an, China, 2024, pp. 412-415."
        ),
        50: (
            "[3] Y. Zhang et al., \"High-Frequency Transmission Characteristic Analysis of TSV-RDL "
            "Interconnects,\" IEEE Transactions on Components, Packaging and Manufacturing "
            "Technology, vol. 14, no. 1, pp. 89-97, Jan. 2024."
        ),
    }

    for index, text in replacements.items():
        replace_by_index(doc, index, text)

    # Keep the paper title readable in the template title block.
    if doc.paragraphs[0].runs:
        doc.paragraphs[0].runs[0].font.size = Pt(18)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
