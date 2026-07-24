from __future__ import annotations

import csv
import json
import re
import subprocess
import sys
import textwrap
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
VERSION_DIR = Path(__file__).resolve().parents[1]
RESULTS_DIR = VERSION_DIR / "results"
TEMPLATE_DOCX = RESULTS_DIR / "Manuscript_template_converted.docx"
OUTPUT_DOCX = VERSION_DIR / "v12建模流程小论文.docx"
ASSET_DIR = VERSION_DIR / "paper_assets"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def metric_row(path: Path, split: str | None = None, device: str | None = None) -> dict[str, str]:
    for row in read_csv_rows(path):
        if split is not None and row.get("split") != split:
            continue
        if device is not None and row.get("device") != device:
            continue
        return row
    raise KeyError(f"No row in {path} for split={split!r}, device={device!r}")


def pct(value: str | float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}%"


def fnum(value: str | float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def convert_template_with_word() -> None:
    if TEMPLATE_DOCX.exists():
        return
    src = VERSION_DIR / "Manuscript.doc"
    script = f"""
$ErrorActionPreference='Stop'
$src='{src}'
$out='{TEMPLATE_DOCX}'
$word=New-Object -ComObject Word.Application
$word.Visible=$false
try {{
  $doc=$word.Documents.Open([ref]$src, [ref]$false, [ref]$true)
  $format=[int]16
  $doc.SaveAs2([ref]$out, [ref]$format)
  $doc.Close([ref]$false)
}} finally {{
  $word.Quit()
}}
"""
    subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True)


def set_font(run, name: str = "Times New Roman", east_asia: str = "SimSun", size: float | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)


def set_paragraph_format(paragraph, *, align=None, before=0, after=3, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "420")


def add_text(paragraph, text: str, *, bold=False, italic=False, size=None):
    run = paragraph.add_run(text)
    set_font(run, size=size)
    run.bold = bold
    run.italic = italic
    return run


def add_paragraph(doc: Document, text: str = "", style: str = "Body Text", *, align=None, before=0, after=3, line=1.05):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, align=align, before=before, after=after, line=line)
    if text:
        add_text(p, text)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=10 if level == 1 else 9)
    run.bold = True
    return p


def add_equation(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        set_font(run, "Cambria Math", "Cambria Math", 8.5)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="BFBFBF", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc: Document, rows: list[list[str]], widths: list[float], caption: str):
    cap = doc.add_paragraph(style="table head")
    set_paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=2)
    caption = re.sub(r"^TABLE\s+[IVXLC]+\.\s*", "", caption)
    add_text(cap, caption, bold=True, size=8)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, "EDEDED")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=7.5)
            run.bold = row_idx == 0
    doc.add_paragraph()
    return table


def add_placeholder_figure(doc: Document, caption: str, note: str, height_cm: float = 4.2):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_borders(table, color="9E9E9E", size="6")
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(height_cm * 10)
    p.paragraph_format.space_after = Pt(height_cm * 10)
    run = p.add_run(note)
    set_font(run, size=8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    add_caption(doc, caption)


def add_caption(doc: Document, caption: str):
    p = doc.add_paragraph(style="figure caption")
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=5, line=1.0)
    caption = re.sub(r"^Fig\.\s*\d+\.\s*", "", caption)
    add_text(p, caption, size=8)


def add_image_figure(doc: Document, image_path: Path, caption: str, width_in: float = 3.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    add_caption(doc, caption)


def download_asset(url: str, name: str) -> Path | None:
    ASSET_DIR.mkdir(exist_ok=True)
    out = ASSET_DIR / name
    if out.exists():
        return out
    try:
        with urllib.request.urlopen(url, timeout=12) as response:
            out.write_bytes(response.read())
        return out
    except Exception:
        return None


def add_references(doc: Document):
    refs = [
        "Z.-X. Ye et al., \"Frequency-domain modeling of interconnects based on assemble neural network for 3-D integration,\" IEEE Trans. Comput.-Aided Design Integr. Circuits Syst., vol. 45, no. 5, pp. 2208-2221, 2026.",
        "H.-S. Yin et al., \"A SPICE-compatible model for on-chip coplanar coupled waveguides up to 110 GHz,\" in Proc. IEEE 7th ICEICT, Xi'an, China, 2024, pp. 412-415.",
        "Prior TSV equivalent-circuit and frequency-domain interconnect modeling literature used in the local v12 workflow notes and extraction scripts.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="Body Text")
        set_paragraph_format(p, before=0, after=2, line=1.0)
        add_text(p, f"[{i}] {ref}", size=7.5)


def build_document() -> None:
    convert_template_with_word()
    doc = Document(TEMPLATE_DOCX)
    clear_document_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    set_two_columns(section)

    single_summary = VERSION_DIR / "results/single_device_paper_nmse_recalculation/single_device_paper_nmse_summary.csv"
    paper_summary = VERSION_DIR / "results/current_best_model/paper_nmse_summary.csv"
    sparam_summary = VERSION_DIR / "results/current_best_model/v08_sparam_summary.csv"
    rdl_val = metric_row(single_summary, split="val", device="RDL")
    tsv_val = metric_row(single_summary, split="val", device="TSV")
    full_test = metric_row(paper_summary, split="test")
    full_train = metric_row(paper_summary, split="train")
    sparam_test = metric_row(sparam_summary, split="test")

    title = doc.add_paragraph(style="paper title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("HFSS-Based Equivalent-Circuit Modeling of Long RDL-TSV Interconnect Chains With a Symmetric Multi-Head Connection Network")
    set_font(run, size=14)
    run.bold = True

    authors = [
        "Author information to be added",
        "Affiliation to be added",
        "E-mail to be added",
    ]
    for line in authors:
        p = doc.add_paragraph(style="Author")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, before=0, after=1, line=1.0)
        add_text(p, line, size=9)

    p = doc.add_paragraph(style="Abstract")
    set_paragraph_format(p, before=8, after=3, line=1.0)
    add_text(
        p,
        "Abstract-Accurate and compact frequency-domain modeling is essential for RDL-TSV interconnect chains in three-dimensional integrated systems, where repeated device sections and interface discontinuities jointly determine the high-frequency response. This paper presents an HFSS-based equivalent-circuit modeling method that combines single-device circuit extraction with a symmetric multi-head pi-type connection network. Equivalent-circuit models are first established for RDL and TSV sections, and multilayer perceptrons are trained to map physical geometry to circuit parameters. A seven-parameter pi network is then inserted at each RDL/TSV interface. To reduce ambiguity under limited training data, the connection parameters are first optimized in a shared form and subsequently expanded into six mirrored heads for the twelve connection positions. The final model is fine-tuned using S11/S21 magnitude and wrapped phase, while the reported paper-style NMSE is evaluated on linear Re/Im(S11,S21) curves. With the selected v12 round-3 checkpoint, the proposed model achieves a test-set average NMSE of "
        + pct(full_test["v08_nn_nmse_mean_percent"], 2)
        + ", compared with "
        + pct(full_test["direct_nmse_mean_percent"], 2)
        + " for the direct cascade without connection correction. The results show that the proposed model substantially improves long-chain prediction accuracy while retaining circuit interpretability and SPICE-compatible structure.",
        size=8,
    )

    p = doc.add_paragraph(style="Keywords")
    set_paragraph_format(p, before=2, after=6, line=1.0)
    add_text(p, "Keywords-RDL; TSV; HFSS; equivalent circuit; S-parameters; pi-type connection network; multi-head neural network; magnitude-phase loss", size=8)

    add_heading(doc, "I. Introduction", 1)
    add_paragraph(
        doc,
        "Three-dimensional (3-D) integration enables high-density heterogeneous systems by vertically stacking dies and connecting them through redistribution layers (RDLs) and through-silicon vias (TSVs). As the interconnect chain becomes longer, the overall frequency response is no longer determined by isolated RDL or TSV sections only. Local discontinuities at RDL/TSV transitions, repeated cascade errors, and phase accumulation jointly affect insertion loss, return loss, and group-delay behavior. Full-wave HFSS simulation remains the most reliable way to characterize these effects, but repeated full-structure simulation is too expensive for design-space exploration, model library construction, and circuit-level co-simulation.",
    )
    add_paragraph(
        doc,
        "A common surrogate strategy is to train a neural network that directly predicts S-parameter curves from geometrical variables. Although this approach can be accurate within a sampled design range, it provides limited circuit interpretability and is difficult to reuse in SPICE-based design flows. Another strategy is to build compact models for individual RDL and TSV devices and then cascade them. This decomposition is more interpretable, but it neglects transition parasitics unless an additional interface model is introduced. For long RDL-TSV chains, even small interface mismatches can accumulate into large full-chain errors.",
    )
    add_paragraph(
        doc,
        "To address this problem, this work introduces a compact modeling flow in which HFSS-derived equivalent-circuit models are used for the RDL and TSV sections, while a seven-parameter pi network is inserted at each interface. The main contributions are as follows. First, SPICE-compatible single-device models are constructed from HFSS data for RDL and TSV sections. Second, a physically interpretable pi-type connection network is used to represent local transition parasitics. Third, a shared-optimization and symmetric multi-head training strategy is proposed to stabilize twelve interface models under a small full-chain dataset. Finally, S11/S21 magnitude and wrapped phase are used for end-to-end fine-tuning, and the final accuracy is reported using the paper-style NMSE on Re/Im(S11,S21).",
    )

    add_heading(doc, "II. Method", 1)
    add_heading(doc, "A. Overall Modeling Flow", 2)
    add_paragraph(
        doc,
        "Fig. 1 summarizes the v12 modeling flow. The RDL and TSV single-device HFSS data in LHS400_Connection2 are first converted into equivalent-circuit parameters, and multilayer perceptrons (MLPs) are trained to map geometrical variables to circuit parameters. A thirteen-section base chain is then assembled in the order TMRDL-TSV-BSMRDL-TSV-TMRDL-TSV-BSMRDL-TSV-TMRDL-TSV-BSMRDL-TSV-TMRDL. Twelve connection networks are inserted between adjacent devices. The connection network is first optimized as one shared seven-parameter circuit for each sample. The optimized targets are then used to train shared parameter networks, which are finally expanded into six mirrored heads and fine-tuned against the full-chain S-parameter response.",
    )
    add_placeholder_figure(doc, "Fig. 1. Overall workflow of the proposed HFSS-equivalent-circuit and symmetric multi-head connection-network modeling method.", "Placeholder for the modeling workflow figure.")

    add_heading(doc, "B. RDL Equivalent-Circuit Model", 2)
    add_paragraph(
        doc,
        "The RDL section is represented by a transmission-line equivalent circuit. For each HFSS sample, the two-port S-parameters are converted into an ABCD matrix, from which the characteristic impedance Zc and propagation constant gamma are extracted as follows:",
    )
    add_equation(
        doc,
        [
            "[A B; C D] = S_to_ABCD(S, Z0)",
            "Zc = sqrt(B / C)",
            "gamma = acosh((A + D) / 2) / l",
        ],
    )
    add_paragraph(
        doc,
        "The per-unit-length RLGC parameters are then obtained by",
    )
    add_equation(
        doc,
        [
            "R + j*omega*L = gamma * Zc",
            "G + j*omega*C = gamma / Zc",
            "R = Re(gamma*Zc),  L = Im(gamma*Zc)/omega",
            "G = Re(gamma/Zc),  C = Im(gamma/Zc)/omega",
        ],
    )
    rdl_image = download_asset("https://ffzhzh.oss-cn-hangzhou.aliyuncs.com/markdown-images/20260712184205051.png", "rdl_equivalent_circuit.png")
    if rdl_image:
        add_image_figure(doc, rdl_image, "Fig. 2. Equivalent circuit used for the RDL single-device model.", width_in=3.0)
    else:
        add_placeholder_figure(doc, "Fig. 2. Equivalent circuit used for the RDL single-device model.", "Placeholder for the RDL equivalent-circuit figure.")
    add_paragraph(
        doc,
        "After parameter extraction, an MLP is trained with pitch, length, width, and thickness as inputs to predict the RDL circuit parameters. The same generic RDL backend is used for both top-metal RDL (TMRDL) and bottom-metal RDL (BSMRDL), while their own geometrical variables are supplied to the model.",
    )

    add_heading(doc, "C. TSV Equivalent-Circuit Model", 2)
    add_paragraph(
        doc,
        "The TSV single-device model uses an equivalent circuit that includes TSV conductor resistance and inductance, oxide capacitance, silicon-substrate capacitance, and substrate loss resistance. Similar to the RDL extraction, TSV S-parameters are converted to ABCD matrices and used to derive RLGC curves. The circuit values are then determined from the low-frequency, high-frequency, and intermediate-frequency points. The TSV model uses r_tsv, h_tsv, and pitch as geometrical inputs.",
    )
    add_equation(
        doc,
        [
            "R_TSV = f_R(RLGC, omega1, omega2, omega3)",
            "L_TSV = f_L(RLGC, omega1, omega2, omega3)",
            "C_ox, C_si, R_si = f_CG(RLGC, omega1, omega2, omega3)",
        ],
    )
    add_placeholder_figure(doc, "Fig. 3. Equivalent circuit used for the TSV single-device model.", "Placeholder for the TSV equivalent-circuit figure.")

    add_heading(doc, "D. Seven-Parameter Pi Connection Network", 2)
    add_paragraph(
        doc,
        "The connection network adopts the v08 seven-parameter pi-type circuit. It contains Cn1, Rn1, Cn2, Rn2, Cn3, Rn3, and Ln1. The two shunt branches describe local capacitive coupling and loss at the two sides of the transition, while the bridging branch captures frequency-dependent interface dispersion. The admittances are defined by",
    )
    add_equation(
        doc,
        [
            "Y1 = j*omega*Cn1 + 1/Rn1",
            "Y2 = j*omega*Cn2 + 1/Rn2",
            "Y3 = j*omega*Cn3 + 1/(Rn3 + j*omega*Ln1)",
        ],
    )
    add_paragraph(doc, "The corresponding ABCD matrix is")
    add_equation(
        doc,
        [
            "A = 1 + Y2/Y3,  B = 1/Y3",
            "C = Y1 + Y2 + Y1*Y2/Y3,  D = 1 + Y1/Y3",
        ],
    )
    add_placeholder_figure(doc, "Fig. 4. Seven-parameter pi-type connection network inserted between adjacent RDL and TSV devices.", "Placeholder for the seven-parameter pi-type connection network.")

    add_heading(doc, "E. Symmetric Multi-Head Training", 2)
    add_paragraph(
        doc,
        "The full chain contains twelve connection positions. Training an independent seven-parameter network for every position would introduce many degrees of freedom and can lead to non-unique parameter solutions when only 150 full-chain training samples are available. Therefore, a two-stage training strategy is used. In the first stage, only one shared seven-parameter circuit is optimized for each sample and repeated at the twelve positions. This stage verifies whether a compact connection correction can reduce the direct-cascade error. In the second stage, seven scalar networks are trained with the architecture 9 -> 30 -> 30 -> 20 -> 1. Their final layers are then expanded into six learned heads, and the heads are mirrored to the twelve connection positions in the order 1,2,3,4,5,6,6,5,4,3,2,1.",
    )
    add_paragraph(
        doc,
        "During multi-head fine-tuning, the loss is computed from S11/S21 magnitude and wrapped phase. The phase error is calculated as angle(pred * conj(target)), which avoids artificial 2*pi discontinuities that appear when unwrapped phases are compared directly. For paper-style reporting, the normalized mean squared error (NMSE) is computed on the linear Re/Im(S11,S21) curves:",
    )
    add_equation(doc, ["NMSE = sum((y_true - y_pred)^2) / sum((y_true - mean(y_true))^2)"])
    add_placeholder_figure(doc, "Fig. 5. Symmetric six-head expansion and mirror mapping for the twelve connection positions.", "Placeholder for the symmetric multi-head network structure.")

    add_heading(doc, "III. Results and Discussion", 1)
    add_heading(doc, "A. Dataset and Experimental Setting", 2)
    add_paragraph(
        doc,
        "The single-device models are trained using the RDL and TSV HFSS data under LHS400_Connection2/train, where 320 samples are used for training and 80 samples are used for validation. The full-chain model is trained with 150 samples from LHS150_50_Connection2/train and tested on 50 samples from LHS150_50_Connection2/test. The selected checkpoint is the round-3 all-150 S-parameter continuation model, which is retained as the current best because the following round produced a higher test NMSE.",
    )
    add_table(
        doc,
        [
            ["Model", "Split", "Count", "Mean NMSE", "Median NMSE"],
            ["RDL single-device", "val", rdl_val["count"], pct(rdl_val["nmse_percent_mean"], 3), pct(rdl_val["nmse_percent_median"], 3)],
            ["TSV single-device", "val", tsv_val["count"], pct(tsv_val["nmse_percent_mean"], 4), pct(tsv_val["nmse_percent_median"], 4)],
            ["Full-chain direct cascade", "test", full_test["count"], pct(full_test["direct_nmse_mean_percent"], 2), pct(full_test["direct_nmse_median_percent"], 2)],
            ["Full-chain v12 model", "test", full_test["count"], pct(full_test["v08_nn_nmse_mean_percent"], 2), pct(full_test["v08_nn_nmse_median_percent"], 2)],
        ],
        [1.55, 0.65, 0.45, 0.9, 0.9],
        "TABLE I. Summary of paper-style NMSE on Re/Im(S11,S21).",
    )

    add_heading(doc, "B. Single-Device Modeling Accuracy", 2)
    add_paragraph(
        doc,
        "The validation results indicate that both single-device backends can reproduce the HFSS frequency response with high accuracy. The RDL validation mean NMSE is "
        + pct(rdl_val["nmse_percent_mean"], 3)
        + ", while the TSV validation mean NMSE is "
        + pct(tsv_val["nmse_percent_mean"], 4)
        + ". These results suggest that the main full-chain error source is not the lack of single-device fitting capability, but the accumulated mismatch introduced by the repeated RDL/TSV transitions and the long cascade.",
    )
    single_plot = VERSION_DIR / "results/single_device_model_vs_hfss_lhs400/plots/TMRDL/random/TMRDL_dut319_random.png"
    if single_plot.exists():
        add_image_figure(doc, single_plot, "Fig. 6. Representative TMRDL single-device comparison between HFSS and the equivalent-circuit model.", width_in=3.15)
    else:
        add_placeholder_figure(doc, "Fig. 6. Representative single-device comparison between HFSS and the equivalent-circuit model.", "Placeholder for the single-device comparison plot.")

    add_heading(doc, "C. Full-Chain Modeling Accuracy", 2)
    add_paragraph(
        doc,
        "On the complete RDL-TSV chain, the direct cascade of single-device models gives a test-set average NMSE of "
        + pct(full_test["direct_nmse_mean_percent"], 2)
        + ", confirming that the interface parasitics cannot be ignored. After introducing the v08 pi-type connection network and applying symmetric multi-head fine-tuning, the test-set average NMSE decreases to "
        + pct(full_test["v08_nn_nmse_mean_percent"], 2)
        + ", and the median NMSE is "
        + pct(full_test["v08_nn_nmse_median_percent"], 2)
        + ". In terms of the direct MSE metric, the proposed v12 model obtains a test mean MSE of "
        + fnum(sparam_test["v08_nn_mse_mean"], 4)
        + ", which is much lower than "
        + fnum(sparam_test["direct_mse_mean"], 4)
        + " from the direct cascade.",
    )
    add_paragraph(
        doc,
        "The training-set average NMSE is "
        + pct(full_train["v08_nn_nmse_mean_percent"], 2)
        + ", while the test-set average NMSE is "
        + pct(full_test["v08_nn_nmse_mean_percent"], 2)
        + ". The gap indicates that some generalization error remains, but the round-3 checkpoint is preferred because the subsequent round-4 continuation increased the test paper-style NMSE. Therefore, the round-3 model is used as current_best_model for this paper draft.",
    )
    full_plot = VERSION_DIR / "results/hfss_v08_symmetric_multihead_lhs150_50_connection2_all150_sparam_continue_round3/comparison_plots/random_test/LHS150_50_Connection2_test_dut172_comparison.png"
    if full_plot.exists():
        add_image_figure(doc, full_plot, "Fig. 7. Representative full-chain S-parameter comparison on a test sample.", width_in=3.15)
    else:
        add_placeholder_figure(doc, "Fig. 7. Representative full-chain S-parameter comparison on a test sample.", "Placeholder for the full-chain comparison plot.")

    add_heading(doc, "D. Discussion", 2)
    add_paragraph(
        doc,
        "Compared with a black-box network that directly outputs the entire S-parameter curve, the proposed model outputs circuit element values for the interface networks. This makes the model easier to inspect and reuse. For example, an abnormal sample can be traced to the capacitance, resistance, or inductance of a specific connection position, which helps distinguish whether the error originates from the single-device backend, the per-sample connection optimization, or the neural mapping. The six-head mirrored structure also provides an engineering constraint that reflects the approximate symmetry of the long chain, thereby reducing the number of free parameters under a limited dataset.",
    )
    add_paragraph(
        doc,
        "The present flow also has limitations. First, the connection-network supervision is obtained from per-sample optimization. If local optima or non-unique circuit parameter sets exist, the supervised targets may still carry ambiguity into the subsequent neural-network training. Second, the current full-chain accuracy remains lower than the single-device accuracy, which indicates that interface modeling and small-sample generalization are still the dominant bottlenecks. Finally, several method figures should be replaced by formal vector drawings or high-resolution images in a camera-ready version; placeholders and captions are retained in this draft for direct replacement.",
    )

    add_heading(doc, "IV. Conclusion", 1)
    add_paragraph(
        doc,
        "This paper presents an HFSS-based equivalent-circuit modeling method for long RDL-TSV interconnect chains. The method combines RDL and TSV single-device circuit models with a seven-parameter pi-type connection network, and it constructs the full-chain model through shared connection optimization, scalar-network pretraining, and symmetric multi-head S-parameter fine-tuning. The selected v12 model achieves an average paper-style NMSE of "
        + pct(full_test["v08_nn_nmse_mean_percent"], 2)
        + " on 50 test samples, which is substantially lower than the direct cascade without connection correction. Because the final representation consists of SPICE-compatible equivalent circuits and compact neural mappings, the proposed approach is suitable for rapid RDL-TSV design exploration and PDK-oriented interconnect model development.",
    )

    add_heading(doc, "References", 1)
    add_references(doc)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_DOCX)
